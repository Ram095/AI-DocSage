"""
Document Processing Utilities

This module handles extraction of text content from various document formats:
- PDF files
- TXT files  
- DOCX files
- CSV files

All functions are explicit and do not use any LangChain abstractions.
"""

import io
from typing import Tuple, Optional
import PyPDF2
from docx import Document
import pandas as pd


# Maximum file size in bytes (10 MB)
MAX_FILE_SIZE = 10 * 1024 * 1024

# Supported file extensions
SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.csv'}


def validate_file(uploaded_file) -> Tuple[bool, str]:
    """
    Validate uploaded file for type and size.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        Tuple of (is_valid, message)
    """
    if uploaded_file is None:
        return False, "No file uploaded"
    
    # Check file size
    file_size = uploaded_file.size
    if file_size > MAX_FILE_SIZE:
        return False, f"File too large. Maximum size is {MAX_FILE_SIZE / (1024*1024):.1f} MB"
    
    # Check file extension
    filename = uploaded_file.name.lower()
    file_ext = '.' + filename.split('.')[-1] if '.' in filename else ''
    
    if file_ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file type. Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
    
    return True, "File validated successfully"


def get_file_metadata(uploaded_file) -> dict:
    """
    Extract metadata from uploaded file.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        Dictionary containing file metadata
    """
    return {
        'filename': uploaded_file.name,
        'size_bytes': uploaded_file.size,
        'size_formatted': format_file_size(uploaded_file.size),
        'type': uploaded_file.type,
        'extension': '.' + uploaded_file.name.split('.')[-1].lower() if '.' in uploaded_file.name else ''
    }


def format_file_size(size_bytes: int) -> str:
    """Format file size in human-readable format."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from PDF file.
    
    Args:
        file_content: Raw bytes of PDF file
        
    Returns:
        Extracted text as string
    """
    text_content = []
    
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
                
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    return '\n\n'.join(text_content)


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text content from TXT file.
    
    Args:
        file_content: Raw bytes of TXT file
        
    Returns:
        Extracted text as string
    """
    try:
        # Try UTF-8 first, then fall back to other encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Unable to decode text file with supported encodings")
    except Exception as e:
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from DOCX file.
    
    Args:
        file_content: Raw bytes of DOCX file
        
    Returns:
        Extracted text as string
    """
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n\n'.join(text_content)
        
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_csv(file_content: bytes) -> str:
    """
    Extract text content from CSV file.
    Converts CSV to a readable text format.
    
    Args:
        file_content: Raw bytes of CSV file
        
    Returns:
        Extracted text as string
    """
    try:
        # Try different encodings
        df = None
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(io.BytesIO(file_content), encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if df is None:
            raise ValueError("Unable to read CSV with supported encodings")
        
        text_parts = []
        
        # Add column info
        text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
        text_parts.append(f"Total rows: {len(df)}")
        text_parts.append("")
        
        # Convert each row to readable text
        for idx, row in df.iterrows():
            row_text = []
            for col in df.columns:
                value = row[col]
                if pd.notna(value):
                    row_text.append(f"{col}: {value}")
            text_parts.append(f"Row {idx + 1}: {' | '.join(row_text)}")
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        raise ValueError(f"Failed to extract text from CSV: {str(e)}")


def extract_text(uploaded_file) -> Tuple[str, dict]:
    """
    Main function to extract text from any supported document type.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        Tuple of (extracted_text, metadata)
        
    Raises:
        ValueError: If file validation fails or text extraction fails
    """
    # Validate file first
    is_valid, message = validate_file(uploaded_file)
    if not is_valid:
        raise ValueError(message)
    
    # Get metadata
    metadata = get_file_metadata(uploaded_file)
    
    # Read file content
    file_content = uploaded_file.read()
    
    # Reset file pointer for potential re-reads
    uploaded_file.seek(0)
    
    # Extract text based on file type
    extension = metadata['extension']
    
    if extension == '.pdf':
        text = extract_text_from_pdf(file_content)
    elif extension == '.txt':
        text = extract_text_from_txt(file_content)
    elif extension == '.docx':
        text = extract_text_from_docx(file_content)
    elif extension == '.csv':
        text = extract_text_from_csv(file_content)
    else:
        raise ValueError(f"Unsupported file type: {extension}")
    
    # Validate extracted text
    if not text or not text.strip():
        raise ValueError("No text content could be extracted from the document")
    
    # Add text stats to metadata
    metadata['char_count'] = len(text)
    metadata['word_count'] = len(text.split())
    
    return text, metadata
